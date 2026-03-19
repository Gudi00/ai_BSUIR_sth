from docx import Document
import os

def create_mock_docs():
    os.makedirs('data/samples', exist_ok=True)
    
    # 1. СТАРАЯ РЕДАКЦИЯ
    doc_old = Document()
    doc_old.add_heading('Договор оказания услуг (Старая редакция)', 0)
    
    doc_old.add_paragraph('1.1. Исполнитель обязан оказать услуги в течение 10 дней.')
    doc_old.add_paragraph('1.2. Заказчик обязуется выплатить 5000 рублей.')
    doc_old.add_paragraph('1.3. Стороны несут ответственность за нарушение обязательств.')
    doc_old.add_paragraph('1.4. Изменение условий допускается по соглашению сторон.')
    
    doc_old.save('data/samples/old_doc.docx')
    
    # 2. НОВАЯ РЕДАКЦИЯ
    doc_new = Document()
    doc_new.add_heading('Договор оказания услуг (Новая редакция)', 0)
    
    # Изменен срок (RED)
    doc_new.add_paragraph('1.1. Исполнитель обязан оказать услуги в течение 15 дней.')
    # Изменена сумма (RED)
    doc_new.add_paragraph('1.2. Заказчик обязуется выплатить 7000 рублей.')
    # Добавлена частица "НЕ" (RED - ОПАСНО!)
    doc_new.add_paragraph('1.3. Стороны НЕ несут ответственность за нарушение обязательств.')
    # Смена модальности (RED)
    doc_new.add_paragraph('1.4. Изменение условий ЗАПРЕЩАЕТСЯ без письменного согласия.')
    # Новый пункт (ADDED)
    doc_new.add_paragraph('1.5. Новый пункт про конфиденциальность.')
    
    doc_new.save('data/samples/new_doc.docx')
    print("Примеры документов созданы в data/samples/")

if __name__ == "__main__":
    create_mock_docs()
