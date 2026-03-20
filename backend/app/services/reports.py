from docx import Document
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
from fpdf import FPDF
import logging
import os
from typing import List
from ..models.block import ComparisonResult

logger = logging.getLogger("uvicorn.error")

def set_cell_background(cell, fill_color: str):
    try:
        shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_color}"/>')
        cell._tc.get_or_add_tcPr().append(shading_elm)
    except Exception as e:
        logger.error(f"DOCX Style Error: {e}")

def generate_docx_report(results: List[ComparisonResult], output_path: str):
    try:
        doc = Document()
        doc.add_heading('Отчет о сравнении документов', 0)
        
        table = doc.add_table(rows=1, cols=2)
        table.style = 'Table Grid'
        
        # Заголовки
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'СТАРАЯ РЕДАКЦИЯ'
        hdr_cells[1].text = 'НОВАЯ РЕДАКЦИЯ / АНАЛИЗ'
        
        COLORS = {"green": "C6EFCE", "yellow": "FFEB9C", "red": "FFC7CE"}

        for res in results:
            row_cells = table.add_row().cells
            
            # Старый текст
            old_val = ""
            if res.old_block:
                old_val = f"{res.old_block.number + ' ' if res.old_block.number else ''}{res.old_block.clean_text}"
            else:
                old_val = "(Блок отсутствовал)"
            row_cells[0].text = old_val
            
            # Новый текст + Анализ
            new_text = ""
            if res.new_block:
                new_text = f"{res.new_block.number + ' ' if res.new_block.number else ''}{res.new_block.clean_text}"
            else:
                new_text = "(Блок удален)"
            
            analysis_parts = []
            if res.change_summary:
                analysis_parts.append(f"Анализ: {res.change_summary}")
            
            if res.risk_triggers:
                trigger_list = ", ".join([f"{t.category}: {t.fragment}" for t in res.risk_triggers])
                analysis_parts.append(f"Триггеры: {trigger_list}")
                
            if analysis_parts:
                new_val = f"{new_text}\n\n" + "\n".join([f"[{p}]" for p in analysis_parts])
            else:
                new_val = new_text
                
            row_cells[1].text = new_val
            
            if res.risk_level in COLORS:
                set_cell_background(row_cells[1], COLORS[res.risk_level])

        doc.save(output_path)
    except Exception as e:
        logger.error(f"CRITICAL DOCX ERROR: {e}")
        raise e

def generate_pdf_report(results: List[ComparisonResult], output_path: str):
    try:
        # Используем fpdf2
        pdf = FPDF()
        
        # Регистрация Unicode шрифта для кириллицы
        # Пытаемся найти системный путь или используем fallback
        font_path = "/usr/share/fonts/TTF/DejaVuSans.ttf"
        if not os.path.exists(font_path):
            font_path = "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf" # Fallback для Ubuntu/Debian

        if os.path.exists(font_path):
            pdf.add_font("DejaVu", "", font_path)
            pdf.add_font("DejaVu", "B", font_path.replace(".ttf", "-Bold.ttf") if "Bold" not in font_path else font_path)
            pdf.set_font("DejaVu", "", 10)
        else:
            logger.warning("Cyrillic font not found, falling back to Helvetica (Russian might not work)")
            pdf.set_font("Helvetica", "", 10)

        pdf.add_page()
        
        pdf.set_font(pdf.font_family, "B", 14)
        pdf.cell(0, 10, 'Отчет о сравнении нормативных актов (Система НАТАША)', ln=True, align='C')
        pdf.set_font(pdf.font_family, "", 10)
        
        COLORS_RGB = {
            "green": (198, 239, 206),
            "yellow": (255, 235, 156),
            "red": (255, 199, 206)
        }

        for res in results:
            old_t = ""
            if res.old_block:
                old_t = f"{res.old_block.number + ' ' if res.old_block.number else ''}{res.old_block.clean_text}"
            else:
                old_t = "(Отсутствует)"
            
            new_t = ""
            if res.new_block:
                new_t = f"{res.new_block.number + ' ' if res.new_block.number else ''}{res.new_block.clean_text}"
            else:
                new_t = "(Удалено)"
            
            # Добавляем анализ
            if res.change_summary:
                new_t += f"\n[AI: {res.change_summary}]"
            if res.human_comment:
                new_t += f"\n[Комментарий: {res.human_comment}]"
            
            y_before = pdf.get_y()
            
            # Рендерим левую колонку (старый текст)
            pdf.multi_cell(90, 5, old_t, border=1)
            y_after_old = pdf.get_y()
            
            # Возвращаемся и рендерим правую колонку (новый текст)
            pdf.set_xy(105, y_before)
            
            if res.risk_level in COLORS_RGB:
                pdf.set_fill_color(*COLORS_RGB[res.risk_level])
                pdf.multi_cell(90, 5, new_t, border=1, fill=True)
            else:
                pdf.multi_cell(90, 5, new_t, border=1)
            
            # Сдвигаемся вниз к самому низкому краю после обеих колонок
            pdf.set_y(max(y_after_old, pdf.get_y()) + 5)
            
            # Если дошли до конца страницы
            if pdf.get_y() > 250:
                pdf.add_page()

        pdf.output(output_path)
    except Exception as e:
        logger.error(f"CRITICAL PDF ERROR: {e}")
        import traceback
        traceback.print_exc()
        raise e
